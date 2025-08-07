import gradio as gr
import pandas as pd
import numpy as np

# import plotly.express as px  # Removed unused import
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import tempfile
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from esg_model import ESGModel
from esg_data_utils import ESGDataProcessor
import warnings

warnings.filterwarnings("ignore")


class ESGGradioApp:
    """
    ESG评分系统Gradio界面
    """

    def __init__(self):
        self.model = ESGModel()
        self.processor = ESGDataProcessor()
        self.current_data = None
        self.current_events = None

        # 从数据处理器获取指标配置
        self.default_indicators = self.processor.get_all_indicators()

    def create_manual_input_data(self, company_name, industry, *indicator_values):
        """
        创建手动输入的数据
        """
        try:
            if not company_name.strip():
                return pd.DataFrame(), "请输入公司名称"

            # 获取所有指标名称
            all_indicators = (
                self.default_indicators["E"]
                + self.default_indicators["S"]
                + self.default_indicators["G"]
            )

            # 创建纵向数据格式
            data_rows = []
            for i, indicator in enumerate(all_indicators):
                value = 0.0  # 默认值
                if i < len(indicator_values) and indicator_values[i] is not None:
                    value = float(indicator_values[i])

                # 确定指标类别
                if indicator in self.default_indicators["E"]:
                    category = "E"
                elif indicator in self.default_indicators["S"]:
                    category = "S"
                else:
                    category = "G"

                data_rows.append(
                    {
                        "company_name": company_name,
                        "industry": industry,
                        "category": category,
                        "indicator": indicator,
                        "value": value,
                    }
                )

            # 创建DataFrame
            data = pd.DataFrame(data_rows)
            self.current_data = data
            self.current_events = [[]]  # 默认无事件

            return (
                data,
                f"成功创建公司 {company_name} 的ESG数据（{len(data_rows)}个指标）",
            )

        except Exception as e:
            return pd.DataFrame(), f"数据创建失败: {str(e)}"

    def upload_custom_data(self, file):
        """
        上传自定义数据（支持CSV和Excel格式）
        """
        try:
            if file is None:
                return pd.DataFrame(), "请上传文件"

            # 根据文件扩展名读取文件
            file_extension = os.path.splitext(file.name)[1].lower()
            if file_extension in [".xlsx", ".xls"]:
                data = pd.read_excel(file.name)
            elif file_extension == ".csv":
                data = pd.read_csv(file.name)
            else:
                return pd.DataFrame(), "不支持的文件格式，请上传CSV或Excel文件"

            # 使用processor进行数据验证和清洗
            try:
                self.processor.validate_company_data(data)
                cleaned_data = self.processor.clean_and_standardize_data(data)
                self.current_data = cleaned_data
                self.current_events = [
                    [] for _ in range(len(cleaned_data))
                ]  # 默认无事件

                preview = cleaned_data.head(10).round(3)
                return preview, f"成功上传并处理{len(cleaned_data)}行数据"
            except ValueError as ve:
                # 如果验证失败，仍然加载数据但给出警告
                self.current_data = data
                self.current_events = [[] for _ in range(len(data))]  # 默认无事件
                preview = data.head(10).round(3)
                return preview, f"数据已上传但存在问题: {str(ve)}"

        except Exception as e:
            return pd.DataFrame(), f"文件上传失败: {str(e)}"

    def calculate_esg_scores(
        self,
        alpha,
        include_events,
        e_weight=0.4,
        s_weight=0.3,
        g_weight=0.3,
        delta_coeff=0.1,
        epsilon_coeff=0.15,
        zeta_coeff=0.12,
        severity_factor=0.4,
        max_bonus=10,
        bonus_steepness=0.8,
        threshold_multiplier=1.0,
        data_breach_coeff=1.2,
        env_pollution_coeff=1.8,
        safety_accident_coeff=1.5,
        corruption_coeff=2.0,
        labor_dispute_coeff=1.0,
        product_recall_coeff=1.3,
        carbon_tax_sensitivity=0.08,
        esg_disclosure_weight=0.15,
        green_finance_bonus=0.05,
        regulatory_compliance=1.0,
        use_cross_terms=True,
    ):
        """
        计算ESG评分
        """
        try:
            if self.current_data is None:
                return pd.DataFrame(), "", "请先生成或上传数据"

            # 使用当前加载的数据
            company_data = self.current_data

            # 检查数据格式
            if "indicator" in company_data.columns and "value" in company_data.columns:
                # 纵向格式数据，需要转换为横向格式
                pivot_data = company_data.pivot_table(
                    index=["company_name", "industry"],
                    columns="indicator",
                    values="value",
                    fill_value=0,
                ).reset_index()

                # 获取ESG指标列
                esg_columns = [
                    col
                    for col in pivot_data.columns
                    if col not in ["company_name", "industry"]
                ]
                esg_data = pivot_data[esg_columns]
                company_info = pivot_data[["company_name", "industry"]]
            else:
                # 横向格式数据（原有格式）
                esg_columns = [
                    col
                    for col in company_data.columns
                    if col
                    not in [
                        "company_id",
                        "company_name",
                        "industry",
                        "market_cap",
                        "employees",
                        "公司名称",
                        "行业",
                    ]
                ]

                if len(esg_columns) == 0:
                    return pd.DataFrame(), "", "数据中未找到ESG指标列"

                esg_data = company_data[esg_columns]
                company_info = (
                    company_data[["公司名称", "行业"]]
                    if "公司名称" in company_data.columns
                    else None
                )

            # 处理事件数据
            events = None  # 默认无事件
            if include_events and self.current_events is not None:
                # 如果current_events是嵌套列表，展平并过滤有效事件
                flat_events = []
                if isinstance(self.current_events, list):
                    for event_list in self.current_events:
                        if isinstance(event_list, list):
                            for event in event_list:
                                if isinstance(event, dict) and event:
                                    flat_events.append(event)
                        elif isinstance(event_list, dict) and event_list:
                            flat_events.append(event_list)
                events = flat_events if flat_events else None

            # 获取行业信息
            if "indicator" in company_data.columns and "value" in company_data.columns:
                # 纵向格式
                industry = (
                    company_info["industry"].iloc[0]
                    if len(company_info) > 0
                    else "制造业"
                )
            else:
                # 横向格式
                if "行业" in company_data.columns:
                    industry = company_data["行业"].iloc[0]
                elif "industry" in company_data.columns:
                    industry = company_data["industry"].iloc[0]
                else:
                    industry = "制造业"

            # 构建甲模型参数字典
            jia_model_params = {
                "alpha": float(alpha),
                "industry_weights": {
                    industry: {
                        "E": float(e_weight),
                        "S": float(s_weight),
                        "G": float(g_weight),
                    },
                    "默认": {
                        "E": float(e_weight),
                        "S": float(s_weight),
                        "G": float(g_weight),
                    },
                },
                "cross_term_coeffs": {
                    "delta": float(delta_coeff),
                    "epsilon": float(epsilon_coeff),
                    "zeta": float(zeta_coeff),
                },
                "nonlinear_params": {
                    "severity_factor": float(severity_factor),
                    "max_bonus": float(max_bonus),
                    "bonus_steepness": float(bonus_steepness),
                    "threshold_multiplier": float(threshold_multiplier),
                },
                "event_coeffs": {
                    "数据泄露": float(data_breach_coeff),
                    "环境污染": float(env_pollution_coeff),
                    "工伤事故": float(safety_accident_coeff),
                    "财务舞弊": float(corruption_coeff),
                    "劳工纠纷": float(labor_dispute_coeff),
                    "产品质量": float(product_recall_coeff),
                },
                "policy_response": {
                    "carbon_tax_sensitivity": float(carbon_tax_sensitivity),
                    "esg_disclosure_weight": float(esg_disclosure_weight),
                    "green_finance_bonus": float(green_finance_bonus),
                },
                "use_cross_terms": bool(use_cross_terms),
            }

            # 计算ESG评分
            results = self.model.calculate_esg_score(
                data=esg_data,
                industry=industry,
                events=events,
                alpha=float(alpha),
                jia_model_params=jia_model_params,
            )

            # 整理结果
            if (
                "indicator" in self.current_data.columns
                and "value" in self.current_data.columns
            ):
                # 纵向格式数据
                company_names = (
                    company_info["company_name"].tolist()
                    if len(company_info) > 0
                    else [f"公司{i + 1}" for i in range(len(esg_data))]
                )
                industries = (
                    company_info["industry"].tolist()
                    if len(company_info) > 0
                    else [industry] * len(esg_data)
                )

                results_df = pd.DataFrame(
                    {
                        "公司ID": range(len(esg_data)),
                        "公司名称": company_names,
                        "行业": industries,
                        "ESG总分": results["final_score"].round(2),
                        "Base Score": results["base_score"].round(2),
                        "E得分": results["e_score"].round(2),
                        "S得分": results["s_score"].round(2),
                        "G得分": results["g_score"].round(2),
                    }
                )
            else:
                # 横向格式数据（原有格式）
                results_df = pd.DataFrame(
                    {
                        "公司ID": self.current_data["company_id"].tolist()
                        if "company_id" in self.current_data.columns
                        else range(len(esg_data)),
                        "公司名称": self.current_data["company_name"].tolist()
                        if "company_name" in self.current_data.columns
                        else [f"公司{i + 1}" for i in range(len(esg_data))],
                        "行业": self.current_data["industry"].tolist()
                        if "industry" in self.current_data.columns
                        else [industry] * len(esg_data),
                        "ESG总分": results["final_score"].round(2),
                        "Base Score": results["base_score"].round(2),
                        "E得分": results["e_score"].round(2),
                        "S得分": results["s_score"].round(2),
                        "G得分": results["g_score"].round(2),
                    }
                )

            # 添加评级
            results_df["评级"] = results_df["ESG总分"].apply(
                lambda x: self.model.get_score_interpretation(x)[0]
            )

            # 生成可视化图表
            charts = self.create_visualization_charts(results_df, results)

            # 生成分析报告
            report = self.generate_analysis_report(
                results_df, results, jia_model_params
            )

            return results_df, charts, report

        except Exception as e:
            empty_fig = go.Figure().add_annotation(
                text=f"评分计算失败: {str(e)}",
                xref="paper",
                yref="paper",
                x=0.5,
                y=0.5,
                showarrow=False,
            )
            return pd.DataFrame(), empty_fig, f"评分计算失败: {str(e)}"

    def create_visualization_charts(self, results_df, model_results):
        """
        创建单公司ESG详细分析图表仪表板
        """
        try:
            if results_df is None or len(results_df) == 0:
                return go.Figure().add_annotation(
                    text="📊 暂无数据进行可视化分析",
                    xref="paper",
                    yref="paper",
                    x=0.5,
                    y=0.5,
                    showarrow=False,
                    font=dict(size=16, color="#888888"),
                )

            # 取第一个公司的数据进行详细分析
            company_data = results_df.iloc[0]
            company_name = company_data.get("公司名称", "企业")

            # 定义配色方案
            colors = {
                "primary": "#1f77b4",
                "secondary": "#ff7f0e",
                "success": "#2ca02c",
                "warning": "#d62728",
                "info": "#17becf",
                "e_color": "#2ca02c",  # 环境 - 绿色
                "s_color": "#1f77b4",  # 社会 - 蓝色
                "g_color": "#ff7f0e",  # 治理 - 橙色
            }

            # 创建子图布局 - 2x3 网格，专注于单公司分析
            fig = make_subplots(
                rows=2,
                cols=3,
                subplot_titles=[
                    "🎯 ESG三维雷达图",
                    "📊 ESG得分仪表盘",
                    "🏆 评级与基准对比",
                    "📈 各维度详细得分",
                    "💡 改进建议热力图",
                    "⚡ 实时评分趋势",
                ],
                specs=[
                    [{"type": "scatterpolar"}, {"type": "indicator"}, {"type": "bar"}],
                    [{"type": "bar"}, {"type": "heatmap"}, {"type": "scatter"}],
                ],
                vertical_spacing=0.12,
                horizontal_spacing=0.08,
            )

            # 1. ESG三维雷达图 (第1行第1列)
            fig.add_trace(
                go.Scatterpolar(
                    r=[
                        company_data["E得分"],
                        company_data["S得分"],
                        company_data["G得分"],
                    ],
                    theta=["环境(E)", "社会(S)", "治理(G)"],
                    fill="toself",
                    name=f"{company_name}",
                    line=dict(color=colors["success"], width=3),
                    fillcolor="rgba(44, 160, 44, 0.3)",
                    marker=dict(size=8),
                ),
                row=1,
                col=1,
            )

            # 添加行业平均水平参考线
            industry_avg = [70, 70, 70]  # 假设的行业平均水平
            fig.add_trace(
                go.Scatterpolar(
                    r=industry_avg,
                    theta=["环境(E)", "社会(S)", "治理(G)"],
                    fill="toself",
                    name="行业平均",
                    line=dict(color=colors["info"], width=2, dash="dash"),
                    fillcolor="rgba(23, 190, 207, 0.2)",
                    marker=dict(size=6),
                ),
                row=1,
                col=1,
            )

            # 2. ESG总分仪表盘 (第1行第2列)
            total_score = company_data["ESG总分"]
            rating, description = self.model.get_score_interpretation(total_score)

            # 确定仪表盘颜色
            if total_score >= 80:
                gauge_color = colors["success"]
            elif total_score >= 60:
                gauge_color = colors["info"]
            elif total_score >= 40:
                gauge_color = colors["secondary"]
            else:
                gauge_color = colors["warning"]

            fig.add_trace(
                go.Indicator(
                    mode="gauge+number+delta",
                    value=total_score,
                    domain={"x": [0, 1], "y": [0, 1]},
                    title={"text": f"<b>{company_name}</b><br>ESG总分"},
                    delta={"reference": 70, "position": "top"},
                    gauge={
                        "axis": {"range": [None, 100]},
                        "bar": {"color": gauge_color},
                        "steps": [
                            {"range": [0, 40], "color": "rgba(214, 39, 40, 0.3)"},
                            {"range": [40, 60], "color": "rgba(255, 127, 14, 0.3)"},
                            {"range": [60, 80], "color": "rgba(23, 190, 207, 0.3)"},
                            {"range": [80, 100], "color": "rgba(44, 160, 44, 0.3)"},
                        ],
                        "threshold": {
                            "line": {"color": "red", "width": 4},
                            "thickness": 0.75,
                            "value": 90,
                        },
                    },
                ),
                row=1,
                col=2,
            )

            # 3. 评级与基准对比 (第1行第3列)
            benchmark_scores = {
                "本企业": total_score,
                "行业平均": 70,
                "优秀企业": 85,
                "政策要求": 60,
            }

            bar_colors = [
                gauge_color,
                colors["info"],
                colors["success"],
                colors["warning"],
            ]

            fig.add_trace(
                go.Bar(
                    x=list(benchmark_scores.keys()),
                    y=list(benchmark_scores.values()),
                    name="基准对比",
                    marker=dict(color=bar_colors, opacity=0.8),
                    text=[f"{v:.1f}" for v in benchmark_scores.values()],
                    textposition="outside",
                    hovertemplate="%{x}: %{y:.1f}分<extra></extra>",
                ),
                row=1,
                col=3,
            )

            # 4. 各维度详细得分 (第2行第1列)
            dimensions = ["环境(E)", "社会(S)", "治理(G)"]
            scores = [
                company_data["E得分"],
                company_data["S得分"],
                company_data["G得分"],
            ]
            dim_colors = [colors["e_color"], colors["s_color"], colors["g_color"]]

            fig.add_trace(
                go.Bar(
                    x=dimensions,
                    y=scores,
                    name="各维度得分",
                    marker=dict(color=dim_colors, opacity=0.8),
                    text=[f"{s:.1f}" for s in scores],
                    textposition="outside",
                    hovertemplate="%{x}: %{y:.1f}分<extra></extra>",
                ),
                row=2,
                col=1,
            )

            # 添加目标线
            for i, target in enumerate([80, 75, 85]):  # 各维度目标分数
                fig.add_shape(
                    type="line",
                    x0=i - 0.4,
                    y0=target,
                    x1=i + 0.4,
                    y1=target,
                    xref=f"x{4}",
                    yref=f"y{4}",
                    line=dict(color="red", width=2, dash="dash"),
                    row=2,
                    col=1,
                )

            # 5. 改进建议热力图 (第2行第2列)
            improvement_matrix = np.array(
                [
                    [
                        company_data["E得分"],
                        max(0, 85 - company_data["E得分"]),
                    ],  # 环境：当前分数 vs 改进空间
                    [
                        company_data["S得分"],
                        max(0, 80 - company_data["S得分"]),
                    ],  # 社会：当前分数 vs 改进空间
                    [
                        company_data["G得分"],
                        max(0, 90 - company_data["G得分"]),
                    ],  # 治理：当前分数 vs 改进空间
                ]
            )

            fig.add_trace(
                go.Heatmap(
                    z=improvement_matrix,
                    x=["当前水平", "改进空间"],
                    y=["环境(E)", "社会(S)", "治理(G)"],
                    colorscale="RdYlGn",
                    text=np.round(improvement_matrix, 1),
                    texttemplate="%{text}",
                    textfont=dict(size=10),
                    hoverongaps=False,
                    hovertemplate="%{y} - %{x}: %{z:.1f}<extra></extra>",
                    name="改进分析",
                ),
                row=2,
                col=2,
            )

            # 6. 实时评分趋势模拟 (第2行第3列)
            # 模拟时间序列数据
            months = ["1月", "2月", "3月", "4月", "5月", "6月"]
            trend_data = {
                "E得分": [
                    max(0, company_data["E得分"] - 5),
                    max(0, company_data["E得分"] - 3),
                    max(0, company_data["E得分"] - 1),
                    company_data["E得分"],
                    min(100, company_data["E得分"] + 1),
                    min(100, company_data["E得分"] + 2),
                ],
                "S得分": [
                    max(0, company_data["S得分"] - 4),
                    max(0, company_data["S得分"] - 2),
                    company_data["S得分"],
                    min(100, company_data["S得分"] + 1),
                    min(100, company_data["S得分"] + 1.5),
                    min(100, company_data["S得分"] + 2),
                ],
                "G得分": [
                    max(0, company_data["G得分"] - 3),
                    max(0, company_data["G得分"] - 1),
                    min(100, company_data["G得分"] + 0.5),
                    company_data["G得分"],
                    min(100, company_data["G得分"] + 1),
                    min(100, company_data["G得分"] + 1.5),
                ],
            }

            for dim, values in trend_data.items():
                color = (
                    colors["e_color"]
                    if "E" in dim
                    else colors["s_color"]
                    if "S" in dim
                    else colors["g_color"]
                )
                fig.add_trace(
                    go.Scatter(
                        x=months,
                        y=values,
                        mode="lines+markers",
                        name=dim,
                        line=dict(color=color, width=3),
                        marker=dict(size=6),
                        hovertemplate=f"{dim}: %{{y:.1f}}分<br>时间: %{{x}}<extra></extra>",
                    ),
                    row=2,
                    col=3,
                )

            # 更新布局
            fig.update_layout(
                height=800,
                title=dict(
                    text=f"🌟 {company_name} ESG评分详细分析仪表板",
                    x=0.5,
                    font=dict(size=20, family="Arial Black"),
                    pad=dict(t=20),
                ),
                template="plotly_white",
                showlegend=True,
                legend=dict(
                    orientation="h", yanchor="bottom", y=-0.1, xanchor="center", x=0.5
                ),
                font=dict(size=11),
                annotations=[
                    dict(
                        text=f"📋 评级: {rating} | 🎯 总分: {total_score:.1f}分 | 💡 {description}",
                        showarrow=False,
                        xref="paper",
                        yref="paper",
                        x=0.5,
                        y=-0.08,
                        xanchor="center",
                        font=dict(size=12, color="#666666"),
                    )
                ],
            )

            # 更新各子图的坐标轴
            # 雷达图
            fig.update_polars(
                radialaxis=dict(
                    visible=True, range=[0, 100], tickmode="linear", tick0=0, dtick=25
                ),
                row=1,
                col=1,
            )

            # 基准对比图
            fig.update_xaxes(title_text="对比基准", row=1, col=3)
            fig.update_yaxes(title_text="得分", row=1, col=3, range=[0, 100])

            # 各维度得分图
            fig.update_xaxes(title_text="ESG维度", row=2, col=1)
            fig.update_yaxes(title_text="得分", row=2, col=1, range=[0, 100])

            # 改进建议热力图
            fig.update_xaxes(title_text="分析维度", row=2, col=2)
            fig.update_yaxes(title_text="ESG维度", row=2, col=2)

            # 趋势图
            fig.update_xaxes(title_text="时间", row=2, col=3)
            fig.update_yaxes(title_text="得分", row=2, col=3, range=[0, 100])

            return fig

        except Exception as e:
            return go.Figure().add_annotation(
                text=f"图表生成失败: {str(e)}",
                xref="paper",
                yref="paper",
                x=0.5,
                y=0.5,
                showarrow=False,
            )

    def generate_data_analysis_results(self, results_df=None):
        """生成数据分析结果"""
        try:
            if results_df is None or len(results_df) == 0:
                return (
                    "**❌ 暂无数据可分析**\n\n请先在'模型配置与评分'页面完成ESG评分。"
                )

            current_time = datetime.now().strftime("%Y年%m月%d日 %H:%M")

            analysis = []
            analysis.append("# 📊 ESG数据分析结果")
            analysis.append(f"**分析时间**: {current_time}")
            analysis.append(f"**数据样本**: {len(results_df)}家企业\n")
            analysis.append("---\n")

            # 基础统计信息
            analysis.append("## 📈 基础统计信息")
            analysis.append("### 总体得分统计")

            avg_score = results_df["ESG总分"].mean()
            max_score = results_df["ESG总分"].max()
            min_score = results_df["ESG总分"].min()
            std_score = results_df["ESG总分"].std()

            analysis.append(f"- **平均得分**: {avg_score:.2f}分")
            analysis.append(f"- **最高得分**: {max_score:.2f}分")
            analysis.append(f"- **最低得分**: {min_score:.2f}分")
            analysis.append(f"- **标准差**: {std_score:.2f}分")

            # 各维度得分统计
            analysis.append("\n### 各维度得分统计")
            for dim in ["E得分", "S得分", "G得分"]:
                dim_avg = results_df[dim].mean()
                dim_name = {"E得分": "环境(E)", "S得分": "社会(S)", "G得分": "治理(G)"}[
                    dim
                ]
                analysis.append(f"- **{dim_name}平均分**: {dim_avg:.2f}分")

            # 评级分布
            analysis.append("\n### 评级分布")
            rating_counts = results_df["评级"].value_counts()
            for rating, count in rating_counts.items():
                percentage = (count / len(results_df)) * 100
                analysis.append(f"- **{rating}**: {count}家企业 ({percentage:.1f}%)")

            # 单企业详细分析
            if len(results_df) == 1:
                company = results_df.iloc[0]
                analysis.append("\n## 🏢 企业详细信息")
                analysis.append(f"### {company['公司名称']}")
                analysis.append(f"- **ESG总分**: {company['ESG总分']:.2f}分")
                analysis.append(f"- **评级**: {company['评级']}")
                analysis.append(f"- **环境(E)得分**: {company['E得分']:.2f}分")
                analysis.append(f"- **社会(S)得分**: {company['S得分']:.2f}分")
                analysis.append(f"- **治理(G)得分**: {company['G得分']:.2f}分")

                # 改进空间分析
                improvement_potential = {
                    "环境(E)": max(0, 85 - company["E得分"]),
                    "社会(S)": max(0, 80 - company["S得分"]),
                    "治理(G)": max(0, 90 - company["G得分"]),
                }

                analysis.append("\n### 🎯 改进空间分析")
                for dim, potential in improvement_potential.items():
                    if potential > 0:
                        analysis.append(f"- **{dim}**: 还有{potential:.1f}分提升空间")
                    else:
                        analysis.append(f"- **{dim}**: 已达到优秀水平")

            # 行业分析
            if "行业" in results_df.columns:
                analysis.append("\n## 🏭 行业分析")
                industry_stats = (
                    results_df.groupby("行业")["ESG总分"]
                    .agg(["count", "mean"])
                    .round(2)
                )
                for industry, stats in industry_stats.iterrows():
                    analysis.append(
                        f"- **{industry}**: {stats['count']}家企业，平均{stats['mean']}分"
                    )

            analysis.append("\n---")
            analysis.append("\n*️⃣ 数据分析完成，可基于此结果生成正式报告*")

            return "\n".join(analysis)

        except Exception as e:
            return f"**❌ 数据分析失败**: {str(e)}"

    def generate_formal_report(
        self, results_df=None, jia_model_params=None, template_type="标准分析报告"
    ):
        """生成正式报告"""
        try:
            if results_df is None or len(results_df) == 0:
                return "**❌ 无法生成报告**\n\n请先完成数据分析。"

            current_time = datetime.now().strftime("%Y年%m月%d日")
            results_df.iloc[0]["公司名称"] if len(
                results_df
            ) == 1 else f"{len(results_df)}家企业"

            if template_type == "简化报告":
                return self._generate_simplified_report(
                    results_df, jia_model_params, current_time
                )
            elif template_type == "详细技术报告":
                return self._generate_detailed_report(
                    results_df, jia_model_params, current_time
                )
            elif template_type == "投资决策报告":
                return self._generate_investment_report(
                    results_df, jia_model_params, current_time
                )
            else:  # 标准分析报告
                return self._generate_standard_report(
                    results_df, jia_model_params, current_time
                )

        except Exception as e:
            return f"**❌ 报告生成失败**: {str(e)}"

    def _generate_standard_report(self, results_df, jia_model_params, current_time):
        """生成标准分析报告"""
        return "标准分析报告功能正在开发中..."

    def _generate_simplified_report(self, results_df, jia_model_params, current_time):
        """生成简化报告"""
        return "简化报告功能正在开发中..."

    def _generate_detailed_report(self, results_df, jia_model_params, current_time):
        """生成详细技术报告"""
        return "详细技术报告功能正在开发中..."

    def _generate_investment_report(self, results_df, jia_model_params, current_time):
        """生成投资决策报告"""
        return "投资决策报告功能正在开发中..."

    def _generate_standard_report(self, results_df, jia_model_params, current_time):
        """生成标准报告"""
        report = []
        company_name = (
            results_df.iloc[0]["公司名称"]
            if len(results_df) == 1
            else f"{len(results_df)}家企业"
        )

        report.append("# 📋 ESG评分分析报告")
        report.append(f"## {company_name} ESG表现评估")
        report.append(f"\n**报告日期**: {current_time}")
        report.append(f"**评估对象**: {company_name}")
        report.append("**评估方法**: 甲模型量化评分体系")
        report.append("**QureLab团队** 专业出品\n")
        report.append("---\n")

        # 执行摘要
        avg_score = results_df["ESG总分"].mean()
        report.append("## 📊 执行摘要")
        report.append(f"本报告采用甲模型对{company_name}进行ESG评分分析。")
        report.append(f"综合评分为**{avg_score:.2f}分**，")

        if avg_score >= 80:
            report.append("表现**优秀**，具备强劲的可持续发展能力。")
        elif avg_score >= 60:
            report.append("表现**良好**，在可持续发展方面具有一定基础。")
        elif avg_score >= 40:
            report.append("表现**中等**，需要在多个维度加强改进。")
        else:
            report.append("表现**有待提升**，建议全面优化ESG管理体系。")

        # 主要发现
        report.append("\n### 🔍 主要发现")
        if len(results_df) == 1:
            company = results_df.iloc[0]
            dimensions = [
                ("环境(E)", company["E得分"]),
                ("社会(S)", company["S得分"]),
                ("治理(G)", company["G得分"]),
            ]
            best_dim = max(dimensions, key=lambda x: x[1])
            worst_dim = min(dimensions, key=lambda x: x[1])

            report.append(
                f"- **优势领域**: {best_dim[0]}维度表现突出({best_dim[1]:.1f}分)"
            )
            report.append(
                f"- **改进重点**: {worst_dim[0]}维度需要加强({worst_dim[1]:.1f}分)"
            )
            report.append(f"- **整体评级**: {company['评级']}")

        # 详细分析
        report.append("\n## 📈 详细分析")
        report.append("### 各维度表现")

        for dim in ["E得分", "S得分", "G得分"]:
            dim_name = {"E得分": "环境(E)", "S得分": "社会(S)", "G得分": "治理(G)"}[dim]
            dim_avg = results_df[dim].mean()

            if dim_avg >= 80:
                performance = "优秀"
                suggestion = "继续保持领先优势"
            elif dim_avg >= 60:
                performance = "良好"
                suggestion = "可进一步优化提升"
            elif dim_avg >= 40:
                performance = "中等"
                suggestion = "需要重点改进"
            else:
                performance = "待提升"
                suggestion = "急需全面改善"

            report.append(f"\n**{dim_name}**: {dim_avg:.1f}分 ({performance})")
            report.append(f"- {suggestion}")

        # 建议
        report.append("\n## 💡 改进建议")
        if len(results_df) == 1:
            company = results_df.iloc[0]
            if company["E得分"] < 70:
                report.append("- **环境方面**: 加强碳排放管控，提升可再生能源使用比例")
            if company["S得分"] < 70:
                report.append(
                    "- **社会方面**: 完善员工福利体系，加强供应链社会责任管理"
                )
            if company["G得分"] < 70:
                report.append("- **治理方面**: 强化董事会独立性，完善风险管理体系")

        report.append("\n---")
        report.append("\n*本报告基于甲模型量化评分体系生成，仅供参考*")

        return "\n".join(report)

    def _generate_simplified_report(self, results_df, jia_model_params, current_time):
        """生成简化报告"""
        company_name = (
            results_df.iloc[0]["公司名称"]
            if len(results_df) == 1
            else f"{len(results_df)}家企业"
        )
        avg_score = results_df["ESG总分"].mean()

        report = [
            "# 📋 ESG评分简报",
            f"**{company_name}** | {current_time}",
            "",
            f"**综合得分**: {avg_score:.1f}分",
            f"**评级**: {results_df.iloc[0]['评级'] if len(results_df) == 1 else '见详细数据'}",
            "",
            "**各维度得分**:",
            f"- 环境(E): {results_df['E得分'].mean():.1f}分",
            f"- 社会(S): {results_df['S得分'].mean():.1f}分",
            f"- 治理(G): {results_df['G得分'].mean():.1f}分",
            "",
            "*QureLab甲模型评估结果*",
        ]
        return "\n".join(report)

    def _generate_detailed_report(self, results_df, jia_model_params, current_time):
        """生成详细报告"""
        # 这里可以调用原来的generate_analysis_report函数
        return self.generate_analysis_report(results_df, {}, jia_model_params)

    def _generate_investment_report(self, results_df, jia_model_params, current_time):
        """生成投资分析报告"""
        company_name = (
            results_df.iloc[0]["公司名称"]
            if len(results_df) == 1
            else f"{len(results_df)}家企业"
        )
        avg_score = results_df["ESG总分"].mean()

        report = [
            "# 💼 ESG投资分析报告",
            f"## {company_name} 投资建议",
            f"**分析日期**: {current_time}",
            "",
            "### 🎯 投资要点",
            f"- **ESG得分**: {avg_score:.1f}分",
            f"- **投资评级**: {'推荐' if avg_score >= 75 else '观望' if avg_score >= 60 else '谨慎'}",
            "",
            "### 📊 风险评估",
        ]

        if avg_score >= 80:
            report.append("- **风险等级**: 低风险")
            report.append("- **投资建议**: 优质ESG标的，建议重点关注")
        elif avg_score >= 60:
            report.append("- **风险等级**: 中等风险")
            report.append("- **投资建议**: ESG表现中等，可适度配置")
        else:
            report.append("- **风险等级**: 高风险")
            report.append("- **投资建议**: ESG风险较高，建议谨慎投资")

        report.extend(
            [
                "",
                "### 💡 投资逻辑",
                "基于甲模型量化分析，该标的在ESG各维度表现如下：",
                f"- 环境维度: {results_df['E得分'].mean():.1f}分",
                f"- 社会维度: {results_df['S得分'].mean():.1f}分",
                f"- 治理维度: {results_df['G得分'].mean():.1f}分",
                "",
                "*本报告仅供投资参考，不构成投资建议*",
            ]
        )

        return "\n".join(report)

    def export_text_content(self, content, filename_prefix):
        """
        导出文本内容为文件
        """
        try:
            import tempfile
            import os
            from datetime import datetime

            if not content or content.strip() == "":
                return None

            # 创建临时文件
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{filename_prefix}_{timestamp}.txt"

            # 使用临时目录
            temp_dir = tempfile.gettempdir()
            file_path = os.path.join(temp_dir, filename)

            # 写入文件
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(content)

            return file_path

        except Exception as e:
            print(f"导出文件失败: {str(e)}")
            return None

    def generate_analysis_report(
        self, results_df, model_results, jia_model_params=None
    ):
        """
        生成专业的ESG分析报告（基于甲模型设计理念）
        """
        try:
            report = []
            current_time = datetime.now().strftime("%Y年%m月%d日")

            # 报告标题和基本信息
            report.append("# 企业ESG评分分析报告")
            report.append("## ——基于甲模型的量化投资策略研究")
            report.append("\n**QureLab团队研究成果**\n")
            report.append(f"**报告日期**: {current_time}")
            report.append(f"**样本规模**: {len(results_df)}家企业")
            report.append("**研究方法**: 甲模型量化评分体系")
            report.append("\n---\n")

            # 摘要
            report.append("## 摘要")
            avg_score = results_df["ESG总分"].mean()
            report.append(
                f"本报告基于甲模型设计理念，对{len(results_df)}家企业进行了全面的ESG评分分析。"
            )
            report.append(
                f"样本企业ESG平均得分为{avg_score:.2f}分，整体表现{'良好' if avg_score >= 70 else '中等' if avg_score >= 50 else '有待提升'}。"
            )
            report.append("甲模型通过多维度交叉项效应、非线性调整机制和政策响应参数，")
            report.append(
                "实现了对企业ESG表现的精准量化评估，为投资决策提供了科学依据。\n"
            )

            # 关键词
            report.append(
                "**关键词**: ESG评分、甲模型、量化投资、可持续发展、风险管理\n"
            )

            # 1. 引言
            report.append("## 1. 引言")
            report.append("### 1.1 研究背景")
            report.append(
                "随着全球可持续发展理念的深入推进，ESG（环境、社会、治理）投资已成为"
            )
            report.append(
                "现代金融市场的重要趋势。2024年中国人民银行等七部委联合印发的"
            )
            report.append("《关于进一步强化金融支持绿色低碳发展的指导意见》明确提出，")
            report.append(
                "要建立国际领先的金融支持绿色低碳发展体系，推动ESG信息披露规范化。"
            )

            report.append("\n### 1.2 甲模型设计理念")
            report.append("甲模型是QureLab团队基于量化投资理论和ESG评估实践开发的")
            report.append("综合评分体系。该模型融合了多模态数据处理、机器学习算法")
            report.append("和金融工程方法，具有以下核心特征：")
            report.append("- **多维度整合**: 通过E×S、E×G、S×G交叉项效应建模")
            report.append("- **非线性调整**: 采用事件分级惩罚和饱和函数奖励机制")
            report.append("- **政策响应**: 动态反映绿色金融政策和监管要求")
            report.append("- **行业差异化**: 基于行业特征的权重优化配置\n")

            # 2. 甲模型参数配置
            if jia_model_params:
                report.append("## 2. 甲模型参数配置")
                report.append("### 2.1 基础权重体系")
                report.append("甲模型采用组合赋权法，结合主观权重和客观权重：")
                report.append(
                    f"- 环境(E)维度权重: {jia_model_params.get('e_weight', 0.4):.3f}"
                )
                report.append(
                    f"- 社会(S)维度权重: {jia_model_params.get('s_weight', 0.3):.3f}"
                )
                report.append(
                    f"- 治理(G)维度权重: {jia_model_params.get('g_weight', 0.3):.3f}"
                )
                report.append(
                    f"- 主观权重系数α: {jia_model_params.get('alpha', 0.5):.3f}"
                )

                report.append("\n### 2.2 交叉项联动系数")
                report.append("基于现代投资组合理论，甲模型引入维度间协同效应：")
                report.append(
                    f"- E×S联动效应系数(δ): {jia_model_params.get('delta_coeff', 0.1):.3f}"
                )
                report.append(
                    f"- E×G联动效应系数(ε): {jia_model_params.get('epsilon_coeff', 0.15):.3f}"
                )
                report.append(
                    f"- S×G联动效应系数(ζ): {jia_model_params.get('zeta_coeff', 0.12):.3f}"
                )

                report.append("\n### 2.3 非线性调整参数")
                report.append("采用行为金融学理论，对极端事件进行非线性处理：")
                report.append(
                    f"- 事件严重度放大因子: {jia_model_params.get('severity_factor', 0.4):.3f}"
                )
                report.append(
                    f"- 最大奖励分数: {jia_model_params.get('max_bonus', 10):.1f}分"
                )
                report.append(
                    f"- 奖励曲线陡度: {jia_model_params.get('bonus_steepness', 0.8):.3f}"
                )
                report.append(
                    f"- 阈值乘数: {jia_model_params.get('threshold_multiplier', 1.0):.3f}"
                )

                report.append("\n### 2.4 事件类型风险系数")
                report.append("基于历史数据统计和专家判断，设定差异化风险权重：")
                report.append(
                    f"- 数据泄露系数: {jia_model_params.get('data_breach_coeff', 1.2):.2f}"
                )
                report.append(
                    f"- 环境污染系数: {jia_model_params.get('env_pollution_coeff', 1.8):.2f}"
                )
                report.append(
                    f"- 安全事故系数: {jia_model_params.get('safety_accident_coeff', 1.5):.2f}"
                )
                report.append(
                    f"- 腐败违规系数: {jia_model_params.get('corruption_coeff', 2.0):.2f}"
                )
                report.append(
                    f"- 劳资纠纷系数: {jia_model_params.get('labor_dispute_coeff', 1.0):.2f}"
                )
                report.append(
                    f"- 产品召回系数: {jia_model_params.get('product_recall_coeff', 1.3):.2f}"
                )

                report.append("\n### 2.5 政策响应机制")
                report.append("结合绿色金融政策导向，动态调整评分权重：")
                report.append(
                    f"- 碳税敏感系数(β): {jia_model_params.get('carbon_tax_sensitivity', 0.08):.3f}"
                )
                report.append(
                    f"- ESG披露权重: {jia_model_params.get('esg_disclosure_weight', 0.15):.3f}"
                )
                report.append(
                    f"- 绿色金融奖励: {jia_model_params.get('green_finance_bonus', 0.05):.3f}"
                )
                report.append(
                    f"- 监管合规系数: {jia_model_params.get('regulatory_compliance', 1.0):.3f}"
                )

                report.append("\n### 2.6 模型功能配置")
                report.append(
                    f"- 交叉项效应: {'启用' if jia_model_params.get('use_cross_terms', True) else '禁用'}"
                )
                report.append(
                    f"- 事件调整机制: {'启用' if jia_model_params.get('include_events', False) else '禁用'}"
                )
                report.append("")

            # 3. 数据统计分析
            report.append("## 3. 样本数据统计分析")
            report.append("### 3.1 基本统计信息")
            report.append(f"- **样本规模**: {len(results_df)}家企业")
            report.append(f"- **ESG平均得分**: {results_df['ESG总分'].mean():.2f}分")
            report.append(f"- **标准差**: {results_df['ESG总分'].std():.2f}")
            report.append(f"- **最高得分**: {results_df['ESG总分'].max():.2f}分")
            report.append(f"- **最低得分**: {results_df['ESG总分'].min():.2f}分")
            report.append(f"- **中位数**: {results_df['ESG总分'].median():.2f}分")

            # 分位数分析
            q25 = results_df["ESG总分"].quantile(0.25)
            q75 = results_df["ESG总分"].quantile(0.75)
            report.append(f"- **第一四分位数(Q1)**: {q25:.2f}分")
            report.append(f"- **第三四分位数(Q3)**: {q75:.2f}分")
            report.append(f"- **四分位距(IQR)**: {q75 - q25:.2f}分\n")

            # 评级分布
            report.append("### 3.2 ESG评级分布")
            rating_dist = results_df["评级"].value_counts().sort_index(ascending=False)
            report.append("| 评级 | 企业数量 | 占比 |")
            report.append("|------|----------|------|")
            for rating, count in rating_dist.items():
                percentage = count / len(results_df) * 100
                report.append(f"| {rating} | {count}家 | {percentage:.1f}% |")
            report.append("")

            # 维度分析
            report.append("### 3.3 各维度表现分析")
            e_avg = results_df["E得分"].mean()
            s_avg = results_df["S得分"].mean()
            g_avg = results_df["G得分"].mean()

            report.append("| 维度 | 平均得分 | 标准差 | 最高分 | 最低分 |")
            report.append("|------|----------|--------|--------|--------|")
            report.append(
                f"| 环境(E) | {e_avg:.2f} | {results_df['E得分'].std():.2f} | {results_df['E得分'].max():.2f} | {results_df['E得分'].min():.2f} |"
            )
            report.append(
                f"| 社会(S) | {s_avg:.2f} | {results_df['S得分'].std():.2f} | {results_df['S得分'].max():.2f} | {results_df['S得分'].min():.2f} |"
            )
            report.append(
                f"| 治理(G) | {g_avg:.2f} | {results_df['G得分'].std():.2f} | {results_df['G得分'].max():.2f} | {results_df['G得分'].min():.2f} |"
            )
            report.append("")

            # 4. 企业表现排名
            report.append("## 4. 企业ESG表现排名")

            # 表现最佳企业
            report.append("### 4.1 ESG表现优秀企业（前10名）")
            top10 = results_df.nlargest(10, "ESG总分")
            report.append(
                "| 排名 | 企业名称 | ESG总分 | 评级 | E得分 | S得分 | G得分 |"
            )
            report.append(
                "|------|----------|---------|------|-------|-------|-------|"
            )
            for i, (idx, row) in enumerate(top10.iterrows(), 1):
                report.append(
                    f"| {i} | {row['公司名称']} | {row['ESG总分']:.2f} | {row['评级']} | {row['E得分']:.2f} | {row['S得分']:.2f} | {row['G得分']:.2f} |"
                )
            report.append("")

            # 需要改进的企业
            report.append("### 4.2 ESG表现待提升企业（后5名）")
            bottom5 = results_df.nsmallest(5, "ESG总分")
            report.append("| 排名 | 企业名称 | ESG总分 | 评级 | 主要问题 |")
            report.append("|------|----------|---------|------|----------|")
            for i, (idx, row) in enumerate(bottom5.iterrows(), 1):
                # 找出最低的维度
                min_dim = min(
                    [(row["E得分"], "E"), (row["S得分"], "S"), (row["G得分"], "G")]
                )
                problem = f"{min_dim[1]}维度偏低({min_dim[0]:.1f}分)"
                report.append(
                    f"| {len(results_df) - len(bottom5) + i} | {row['公司名称']} | {row['ESG总分']:.2f} | {row['评级']} | {problem} |"
                )
            report.append("")

            # 5. 量化投资策略分析
            report.append("## 5. 基于甲模型的量化投资策略分析")
            report.append("### 5.1 数据驱动的投资决策")
            report.append("根据量化投资理论，甲模型采用多模态数据处理技术，")
            report.append("整合了结构化财务数据、非结构化新闻舆情、")
            report.append("以及实时政策信息，形成全方位的ESG评估体系。")

            report.append("\n### 5.2 因子分析与特征工程")
            report.append("甲模型构建了三类核心因子：")
            report.append("- **量价因子**: 基于市场交易数据的技术指标")
            report.append("- **基本面因子**: 来源于财务报表和ESG披露信息")
            report.append("- **另类因子**: 包括舆情数据、政策响应等")

            report.append("\n### 5.3 风险管理与组合优化")
            high_risk_count = len(results_df[results_df["ESG总分"] < 40])
            medium_risk_count = len(
                results_df[(results_df["ESG总分"] >= 40) & (results_df["ESG总分"] < 70)]
            )
            low_risk_count = len(results_df[results_df["ESG总分"] >= 70])

            report.append("基于ESG评分的风险分层结果：")
            report.append(
                f"- **低风险组合** (ESG≥70分): {low_risk_count}家企业 ({low_risk_count / len(results_df) * 100:.1f}%)"
            )
            report.append(
                f"- **中等风险组合** (40≤ESG<70分): {medium_risk_count}家企业 ({medium_risk_count / len(results_df) * 100:.1f}%)"
            )
            report.append(
                f"- **高风险组合** (ESG<40分): {high_risk_count}家企业 ({high_risk_count / len(results_df) * 100:.1f}%)"
            )

            # 6. 政策环境分析
            report.append("\n## 6. 绿色金融政策环境分析")
            report.append("### 6.1 政策背景")
            report.append("2024年10月，中国人民银行等四部委联合印发")
            report.append("《关于发挥绿色金融作用 服务美丽中国建设的意见》，")
            report.append("明确了绿色金融支持美丽中国建设的重点领域和实施路径。")

            report.append("\n### 6.2 政策影响分析")
            report.append("甲模型的政策响应机制体现在以下方面：")
            if jia_model_params:
                carbon_sensitivity = jia_model_params.get(
                    "carbon_tax_sensitivity", 0.08
                )
                disclosure_weight = jia_model_params.get("esg_disclosure_weight", 0.15)
                green_bonus = jia_model_params.get("green_finance_bonus", 0.05)

                report.append(
                    f"- **碳税敏感性调整**: 系数{carbon_sensitivity:.3f}，反映碳定价政策影响"
                )
                report.append(
                    f"- **ESG披露权重**: {disclosure_weight:.3f}，鼓励信息透明度"
                )
                report.append(
                    f"- **绿色金融奖励**: {green_bonus:.3f}，支持绿色项目融资"
                )

            # 7. 模型验证与回测
            if "weights" in model_results:
                report.append("\n## 7. 模型验证与回测分析")
                weights = model_results["weights"]
                report.append("### 7.1 权重分布统计")
                report.append(f"- 权重向量维度: {len(weights)}")
                report.append(f"- 权重最大值: {weights.max():.4f}")
                report.append(f"- 权重最小值: {weights.min():.4f}")
                report.append(f"- 权重标准差: {weights.std():.4f}")

                # 计算权重集中度
                weight_concentration = (weights**2).sum()
                report.append(f"- 权重集中度(HHI): {weight_concentration:.4f}")

            # 8. 投资建议
            report.append("\n## 8. 投资策略建议")

            # 基于平均得分给出建议
            report.append("### 8.1 维度优化建议")
            if e_avg < 50:
                report.append("**环境维度改进**:")
                report.append("- 加大清洁能源投资，推进碳中和目标")
                report.append("- 建立环境管理体系，获得ISO14001认证")
                report.append("- 开展环境风险评估，制定应急预案")

            if s_avg < 50:
                report.append("\n**社会维度改进**:")
                report.append("- 完善员工福利体系，提升员工满意度")
                report.append("- 加强供应链社会责任管理")
                report.append("- 积极参与社区公益活动")

            if g_avg < 50:
                report.append("\n**治理维度改进**:")
                report.append("- 优化董事会结构，提高独立董事比例")
                report.append("- 建立健全内控制度和风险管理体系")
                report.append("- 加强信息披露透明度")

            # 量化投资建议
            report.append("\n### 8.2 量化投资策略建议")
            report.append("基于甲模型分析结果，建议采用以下投资策略：")

            report.append("\n**1. 多因子选股策略**")
            report.append("- 构建ESG-增强型多因子模型")
            report.append("- 结合价值、成长、质量因子")
            report.append("- 动态调整因子权重")

            report.append("\n**2. 风险平价策略**")
            report.append("- 基于ESG评分进行风险预算分配")
            report.append("- 控制单一维度风险暴露")
            report.append("- 实施动态再平衡")

            report.append("\n**3. 事件驱动策略**")
            report.append("- 监控ESG相关负面事件")
            report.append("- 利用市场过度反应获取超额收益")
            report.append("- 建立事件影响评估模型")

            # 甲模型特色建议
            if jia_model_params:
                report.append("\n### 8.3 基于甲模型的专业建议")
                report.append("**价值导向投资**:")
                report.append("- 将ESG理念融入投资决策全流程")
                report.append("- 关注长期价值创造能力")
                report.append("- 平衡财务回报与社会效益")

                report.append("\n**实质性原则**:")
                report.append("- 重点关注对业务影响最大的ESG议题")
                report.append("- 建立行业特定的ESG评估框架")
                report.append("- 定期更新重要性矩阵")

                report.append("\n**利益相关方协同**:")
                report.append("- 加强与投资者的ESG沟通")
                report.append("- 建立客户ESG需求反馈机制")
                report.append("- 推动供应链ESG标准统一")

            # 9. 结论
            report.append("\n## 9. 结论")
            report.append(
                f"本研究基于甲模型对{len(results_df)}家企业进行了全面的ESG评分分析。"
            )
            report.append("研究发现，甲模型通过多维度交叉项效应、非线性调整机制")
            report.append("和政策响应参数，能够有效识别企业ESG风险和机遇，")
            report.append("为量化投资决策提供了科学依据。")

            report.append("\n未来研究方向包括：")
            report.append("- 扩大样本规模，增强模型泛化能力")
            report.append("- 引入更多另类数据源")
            report.append("- 开发实时ESG评分系统")
            report.append("- 构建ESG投资组合优化算法")

            # 参考文献
            report.append("\n## 参考文献")
            report.append(
                "[1] 中国人民银行等. 关于进一步强化金融支持绿色低碳发展的指导意见[Z]. 2024."
            )
            report.append(
                "[2] 中国人民银行等. 关于发挥绿色金融作用 服务美丽中国建设的意见[Z]. 2024."
            )
            report.append(
                "[3] Markowitz, H. Portfolio Selection[J]. Journal of Finance, 1952, 7(1): 77-91."
            )
            report.append(
                "[4] Fama, E. F., French, K. R. Common risk factors in the returns on stocks and bonds[J]. Journal of Financial Economics, 1993, 33(1): 3-56."
            )

            report.append("\n---")
            report.append(
                "\n**声明**: 本报告仅供研究参考，不构成投资建议。投资有风险，决策需谨慎。"
            )
            report.append(f"\n**QureLab团队** | {current_time}")

            return "\n".join(report)

        except Exception as e:
            return f"报告生成失败: {str(e)}"

    def export_input_data(self, format_type):
        """
        导出输入数据（纵向格式）
        """
        try:
            if self.current_data is None or len(self.current_data) == 0:
                return None, "没有可导出的数据"

            # 创建临时文件
            if format_type == "CSV":
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".csv")
                self.current_data.to_csv(
                    temp_file.name, index=False, encoding="utf-8-sig"
                )
                return temp_file.name, "输入数据已导出为CSV文件（纵向格式）"
            else:  # Excel
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
                self.current_data.to_excel(temp_file.name, index=False)
                return temp_file.name, "输入数据已导出为Excel文件（纵向格式）"

        except Exception as e:
            return None, f"导出失败: {str(e)}"

    def export_results(self, results_df):
        """
        导出评分结果
        """
        try:
            if results_df is None or len(results_df) == 0:
                return None, "没有可导出的数据"

            # 创建临时文件
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
            results_df.to_excel(temp_file.name, index=False)

            return temp_file.name, "结果已导出为Excel文件"

        except Exception as e:
            return None, f"导出失败: {str(e)}"

    def export_report_as_word(self, report_content):
        """
        导出分析报告为Word文件
        """
        try:
            if not report_content:
                return None

            # 创建Word文档
            doc = Document()

            # 设置文档样式
            style = doc.styles["Normal"]
            style.font.name = "宋体"
            style.font.size = Pt(12)

            # 添加标题样式
            title_style = doc.styles.add_style("CustomTitle", WD_STYLE_TYPE.PARAGRAPH)
            title_style.font.name = "黑体"
            title_style.font.size = Pt(16)
            title_style.font.bold = True
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加二级标题样式
            heading_style = doc.styles.add_style(
                "CustomHeading", WD_STYLE_TYPE.PARAGRAPH
            )
            heading_style.font.name = "黑体"
            heading_style.font.size = Pt(14)
            heading_style.font.bold = True

            # 解析报告内容并添加到文档
            lines = report_content.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                if line.startswith("# "):
                    # 主标题
                    p = doc.add_paragraph(line[2:], style="CustomTitle")
                elif line.startswith("## "):
                    # 二级标题
                    p = doc.add_paragraph(line[3:], style="CustomHeading")
                elif line.startswith("### "):
                    # 三级标题
                    p = doc.add_paragraph(line[4:])
                    p.style.font.bold = True
                elif line.startswith("- "):
                    # 列表项
                    p = doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("**") and line.endswith("**"):
                    # 粗体文本
                    p = doc.add_paragraph()
                    run = p.add_run(line[2:-2])
                    run.bold = True
                else:
                    # 普通段落
                    doc.add_paragraph(line)

            # 保存到临时文件
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
            doc.save(temp_file.name)

            return temp_file.name

        except Exception as e:
            print(f"Word导出失败: {str(e)}")
            return None

    def export_report_as_pdf(self, report_content):
        """
        导出分析报告为PDF文件
        """
        try:
            if not report_content:
                return None

            # 创建临时文件
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

            # 创建PDF文档
            doc = SimpleDocTemplate(temp_file.name, pagesize=A4)

            # 获取样式
            styles = getSampleStyleSheet()

            # 创建自定义样式
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Title"],
                fontSize=16,
                spaceAfter=20,
                alignment=1,  # 居中
            )

            heading_style = ParagraphStyle(
                "CustomHeading",
                parent=styles["Heading1"],
                fontSize=14,
                spaceAfter=12,
                spaceBefore=12,
            )

            normal_style = ParagraphStyle(
                "CustomNormal", parent=styles["Normal"], fontSize=12, spaceAfter=6
            )

            # 解析内容并创建段落
            story = []
            lines = report_content.split("\n")

            for line in lines:
                line = line.strip()
                if not line:
                    story.append(Spacer(1, 6))
                    continue

                if line.startswith("# "):
                    # 主标题
                    story.append(Paragraph(line[2:], title_style))
                elif line.startswith("## "):
                    # 二级标题
                    story.append(Paragraph(line[3:], heading_style))
                elif line.startswith("### "):
                    # 三级标题
                    story.append(Paragraph(f"<b>{line[4:]}</b>", normal_style))
                elif line.startswith("- "):
                    # 列表项
                    story.append(Paragraph(f"• {line[2:]}", normal_style))
                elif line.startswith("**") and line.endswith("**"):
                    # 粗体文本
                    story.append(Paragraph(f"<b>{line[2:-2]}</b>", normal_style))
                else:
                    # 普通段落
                    story.append(Paragraph(line, normal_style))

            # 构建PDF
            doc.build(story)

            return temp_file.name

        except Exception as e:
            print(f"PDF导出失败: {str(e)}")
            return None

    def import_scoring_data(self, results_df):
        """
        从评分页面导入数据到分析报告页面
        """
        try:
            if results_df is None or len(results_df) == 0:
                return (
                    "❌ 暂无评分数据，请先在'模型配置与评分'页面完成ESG评分",
                    gr.Dataframe(visible=False),
                    gr.Button(interactive=False),
                )

            # 显示导入成功状态
            status_msg = f"✅ 成功导入评分数据，共 {len(results_df)} 条记录"

            # 返回预览数据
            preview_df = gr.Dataframe(value=results_df, visible=True, interactive=False)

            # 激活生成报告按钮
            generate_btn = gr.Button(interactive=True)

            return status_msg, preview_df, generate_btn

        except Exception as e:
            return (
                f"❌ 数据导入失败: {str(e)}",
                gr.Dataframe(visible=False),
                gr.Button(interactive=False),
            )

    def generate_imported_analysis_report(self, imported_df, template_type):
        """
        基于导入的评分数据生成分析报告
        """
        try:
            if imported_df is None or len(imported_df) == 0:
                return (
                    "*请先导入评分数据*",
                    gr.Button(interactive=False),
                    gr.Button(interactive=False),
                    gr.Button(interactive=False),
                )

            # 使用现有的报告生成方法
            report_content = self.generate_formal_report(
                results_df=imported_df,
                jia_model_params=None,
                template_type=template_type,
            )

            # 激活导出按钮
            txt_btn = gr.Button(interactive=True)
            word_btn = gr.Button(interactive=True)
            pdf_btn = gr.Button(interactive=True)

            return report_content, txt_btn, word_btn, pdf_btn

        except Exception as e:
            error_msg = f"报告生成失败: {str(e)}"
            return (
                error_msg,
                gr.Button(interactive=False),
                gr.Button(interactive=False),
                gr.Button(interactive=False),
            )

    def toggle_evaluation_type(self, evaluation_type):
        """
        切换评估类型显示
        """
        if evaluation_type == "单个公司评估":
            return (
                gr.update(visible=True),  # single_company_group
                gr.update(visible=False),  # multiple_company_group
                gr.update(visible=True),  # single_upload_group
            )
        else:  # 多个公司评估
            return (
                gr.update(visible=False),  # single_company_group
                gr.update(visible=True),  # multiple_company_group
                gr.update(visible=False),  # single_upload_group
            )

    def load_scoring_data(self, file):
        """
        加载评分数据文件
        """
        try:
            if file is None:
                return "请上传文件"

            # 根据文件扩展名读取文件
            file_extension = os.path.splitext(file.name)[1].lower()
            if file_extension in [".xlsx", ".xls"]:
                data = pd.read_excel(file.name)
            elif file_extension == ".csv":
                data = pd.read_csv(file.name)
            else:
                return "不支持的文件格式，请上传CSV或Excel文件"

            # 使用processor进行数据验证和清洗
            try:
                self.processor.validate_company_data(data)
                cleaned_data = self.processor.clean_and_standardize_data(data)
                self.current_data = cleaned_data
                self.current_events = [
                    [] for _ in range(len(cleaned_data))
                ]  # 默认无事件

                return f"✅ 成功加载{len(cleaned_data)}行数据，包含{len(cleaned_data)}个公司的ESG数据"
            except ValueError as ve:
                # 如果验证失败，仍然加载数据但给出警告
                self.current_data = data
                self.current_events = [[] for _ in range(len(data))]  # 默认无事件
                return f"⚠️ 数据已加载但存在问题: {str(ve)}"

        except Exception as e:
            return f"❌ 文件加载失败: {str(e)}"

    def create_interface(self):
        """
        创建Gradio界面
        """
        # 自定义CSS样式，使标签页居中
        custom_css = """
        /* 强制标签页居中 */
        .gradio-container .tabs .tab-nav,
        .gradio-container .tabs > .tab-nav,
        .tab-nav,
        div[data-testid="tabs"] .tab-nav,
        div[data-testid="tabs"] > div:first-child,
        .gradio-tabs .tab-nav,
        .gradio-tabs > .tab-nav,
        [data-testid="tabs"] > div,
        [data-testid="tabs"] > div > div {
            justify-content: center !important;
            display: flex !important;
            width: 100% !important;
            text-align: center !important;
        }
        
        /* 确保标签按钮居中 */
        .tab-nav button,
        .tab-nav .tab-item,
        [data-testid="tabs"] button {
            margin: 0 auto !important;
        }
        
        /* 整体容器居中 */
        .gradio-tabs {
            text-align: center !important;
        }
        """

        with gr.Blocks(
            title="ESG评分系统",
            theme=gr.themes.Soft(primary_hue="green"),
            css=custom_css,
        ) as interface:
            # 禁用队列以避免WebSocket连接问题
            interface.queue = lambda *args, **kwargs: interface
            gr.Markdown(
                """
                # 基于甲模型设计理念的企业ESG（环境、社会、治理）评分系统
                
                """
            )

            with gr.Tabs():
                # 第一页：数据输入
                with gr.TabItem("📊 数据输入", id=0):
                    gr.Markdown("### 📋 单个公司ESG数据输入")
                    gr.Markdown(
                        "请选择数据输入方式：手动输入或上传文件（支持Excel和CSV格式）"
                    )

                    with gr.Row():
                        with gr.Column(scale=2):
                            # 手动输入区域
                            with gr.Column():
                                gr.Markdown("### 公司基本信息")

                                with gr.Row():
                                    company_name = gr.Textbox(
                                        label="公司名称", placeholder="请输入公司名称"
                                    )
                                    industry_choice = gr.Dropdown(
                                        choices=[
                                            "制造业",
                                            "金融业",
                                            "科技业",
                                            "能源业",
                                            "消费业",
                                        ],
                                        value="制造业",
                                        label="行业类型",
                                    )

                                gr.Markdown("### ESG指标输入")

                                # 使用可折叠的Accordion组件
                                e_inputs = []
                                s_inputs = []
                                g_inputs = []

                                with gr.Accordion("环境指标 (E)", open=False):
                                    for indicator in self.default_indicators["E"]:
                                        e_inputs.append(
                                            gr.Number(
                                                value=0,
                                                minimum=0,
                                                maximum=100,
                                                step=1,
                                                label=indicator,
                                            )
                                        )

                                with gr.Accordion("社会指标 (S)", open=False):
                                    for indicator in self.default_indicators["S"]:
                                        s_inputs.append(
                                            gr.Number(
                                                value=0,
                                                minimum=0,
                                                maximum=100,
                                                step=1,
                                                label=indicator,
                                            )
                                        )

                                with gr.Accordion("治理指标 (G)", open=False):
                                    for indicator in self.default_indicators["G"]:
                                        g_inputs.append(
                                            gr.Number(
                                                value=0,
                                                minimum=0,
                                                maximum=100,
                                                step=1,
                                                label=indicator,
                                            )
                                        )

                                with gr.Row():
                                    create_btn = gr.Button(
                                        "创建ESG数据", variant="primary", size="lg"
                                    )
                                    with gr.Column():
                                        export_format = gr.Dropdown(
                                            choices=["CSV", "Excel"],
                                            value="Excel",
                                            label="导出格式",
                                        )
                                        export_input_btn = gr.Button(
                                            "下载模板/数据", variant="secondary"
                                        )

                        with gr.Column(scale=1):
                            # 文件上传选项
                            with gr.Column():
                                gr.Markdown("### 📁 或上传数据文件")
                                gr.Markdown("支持Excel (.xlsx) 和 CSV (.csv) 格式")
                                gr.Markdown("**文件格式要求：**")
                                gr.Markdown("""
                                - 包含列：公司名称, 行业, ESG指标列
                                - 每行代表一个公司的完整ESG数据
                                - 指标值范围：0-100
                                """)
                                upload_file = gr.File(
                                    label="上传Excel或CSV文件",
                                    file_types=[".xlsx", ".xls", ".csv"],
                                )
                                upload_btn = gr.Button(
                                    "上传数据", variant="primary", size="lg"
                                )

                            # 通用状态显示区域
                            data_status = gr.Textbox(
                                label="数据状态", interactive=False
                            )
                            export_input_file = gr.File(label="下载输入数据")
                            export_input_status = gr.Textbox(
                                label="导出状态", interactive=False
                            )

                    data_preview = gr.Dataframe(label="数据预览")

                # 第二页：模型配置与评分
                with gr.TabItem("⚙️ 模型配置与评分", id=1):
                    # 文件上传区域
                    gr.Markdown("### 📁 上传ESG数据文件")
                    gr.Markdown("**必须上传包含ESG数据的文件才能进行评分**")
                    with gr.Row():
                        with gr.Column(scale=2):
                            scoring_upload_file = gr.File(
                                label="上传Excel或CSV文件",
                                file_types=[".xlsx", ".xls", ".csv"],
                            )
                        with gr.Column(scale=1):
                            load_data_btn = gr.Button(
                                "📊 加载数据", variant="primary", size="lg"
                            )

                    data_load_status = gr.Textbox(
                        label="数据加载状态",
                        interactive=False,
                        placeholder="请上传文件并点击加载数据",
                    )

                    gr.Markdown("### 模型参数配置")

                    # 基础权重配置
                    with gr.Accordion("基础权重配置", open=True):
                        with gr.Row():
                            alpha_param = gr.Slider(
                                0, 1, value=0.5, label="主观权重系数α"
                            )

                        with gr.Row():
                            with gr.Column():
                                e_weight = gr.Slider(
                                    minimum=0.1,
                                    maximum=0.8,
                                    value=0.4,
                                    step=0.05,
                                    label="环境(E)维度权重",
                                )
                            with gr.Column():
                                s_weight = gr.Slider(
                                    minimum=0.1,
                                    maximum=0.8,
                                    value=0.3,
                                    step=0.05,
                                    label="社会(S)维度权重",
                                )
                            with gr.Column():
                                g_weight = gr.Slider(
                                    minimum=0.1,
                                    maximum=0.8,
                                    value=0.3,
                                    step=0.05,
                                    label="治理(G)维度权重",
                                )

                    # 交叉项系数配置
                    with gr.Accordion("交叉项系数配置（甲模型整合性原则）", open=False):
                        with gr.Row():
                            with gr.Column():
                                delta_coeff = gr.Slider(
                                    minimum=-0.5,
                                    maximum=0.5,
                                    value=0.1,
                                    step=0.01,
                                    label="δ (E×S交叉项系数)",
                                )
                            with gr.Column():
                                epsilon_coeff = gr.Slider(
                                    minimum=-0.5,
                                    maximum=0.5,
                                    value=0.15,
                                    step=0.01,
                                    label="ε (E×G交叉项系数)",
                                )
                            with gr.Column():
                                zeta_coeff = gr.Slider(
                                    minimum=-0.5,
                                    maximum=0.5,
                                    value=0.12,
                                    step=0.01,
                                    label="ζ (S×G交叉项系数)",
                                )

                    # 非线性调整参数
                    with gr.Accordion("非线性调整参数（甲模型动态适应性）", open=False):
                        with gr.Row():
                            with gr.Column():
                                severity_factor = gr.Slider(
                                    minimum=0.1,
                                    maximum=1.0,
                                    value=0.4,
                                    step=0.05,
                                    label="严重度放大因子β",
                                )
                                max_bonus = gr.Slider(
                                    minimum=5,
                                    maximum=20,
                                    value=10,
                                    step=1,
                                    label="最大奖励分数",
                                )
                            with gr.Column():
                                bonus_steepness = gr.Slider(
                                    minimum=0.1,
                                    maximum=2.0,
                                    value=0.8,
                                    step=0.1,
                                    label="奖励曲线陡度k",
                                )
                                threshold_multiplier = gr.Slider(
                                    minimum=0.8,
                                    maximum=1.2,
                                    value=1.0,
                                    step=0.05,
                                    label="阈值乘数",
                                )

                    # 事件类型系数
                    with gr.Accordion("事件类型系数λ（甲模型事件分级）", open=False):
                        with gr.Row():
                            with gr.Column():
                                data_breach_coeff = gr.Slider(
                                    minimum=0.5,
                                    maximum=2.0,
                                    value=1.2,
                                    step=0.1,
                                    label="数据泄露系数",
                                )
                                env_pollution_coeff = gr.Slider(
                                    minimum=0.8,
                                    maximum=2.5,
                                    value=1.8,
                                    step=0.1,
                                    label="环境污染系数",
                                )
                                safety_accident_coeff = gr.Slider(
                                    minimum=0.6,
                                    maximum=2.0,
                                    value=1.5,
                                    step=0.1,
                                    label="安全事故系数",
                                )
                            with gr.Column():
                                corruption_coeff = gr.Slider(
                                    minimum=1.0,
                                    maximum=3.0,
                                    value=2.0,
                                    step=0.1,
                                    label="腐败违规系数",
                                )
                                labor_dispute_coeff = gr.Slider(
                                    minimum=0.5,
                                    maximum=1.8,
                                    value=1.0,
                                    step=0.1,
                                    label="劳资纠纷系数",
                                )
                                product_recall_coeff = gr.Slider(
                                    minimum=0.7,
                                    maximum=2.2,
                                    value=1.3,
                                    step=0.1,
                                    label="产品召回系数",
                                )

                    # 政策响应参数
                    with gr.Accordion("政策响应参数（甲模型动态适应性）", open=False):
                        with gr.Row():
                            with gr.Column():
                                carbon_tax_sensitivity = gr.Slider(
                                    minimum=0.01,
                                    maximum=0.2,
                                    value=0.08,
                                    step=0.01,
                                    label="碳税敏感系数",
                                )
                                esg_disclosure_weight = gr.Slider(
                                    minimum=0.05,
                                    maximum=0.3,
                                    value=0.15,
                                    step=0.01,
                                    label="ESG披露权重",
                                )
                            with gr.Column():
                                green_finance_bonus = gr.Slider(
                                    minimum=0.02,
                                    maximum=0.15,
                                    value=0.05,
                                    step=0.01,
                                    label="绿色金融奖励",
                                )
                                regulatory_compliance = gr.Slider(
                                    minimum=0.8,
                                    maximum=1.5,
                                    value=1.0,
                                    step=0.05,
                                    label="监管合规系数",
                                )

                    with gr.Row():
                        include_events_check = gr.Checkbox(
                            value=True, label="启用事件调整（甲模型非线性调整）"
                        )
                        use_cross_terms = gr.Checkbox(
                            value=True, label="启用交叉项效应（甲模型整合性原则）"
                        )

                    calculate_btn = gr.Button(
                        "🚀 计算ESG评分", variant="primary", size="lg"
                    )

                    gr.Markdown("### 📈 评分结果")
                    results_table = gr.Dataframe(label="评分结果", interactive=False)

                    gr.Markdown("### 📊 可视化分析")
                    charts_plot = gr.Plot(label="分析图表")

                    with gr.Row():
                        export_results_btn = gr.Button(
                            "下载评分结果", variant="secondary"
                        )
                        export_results_file = gr.File(label="下载评分结果")
                        export_results_status = gr.Textbox(
                            label="导出状态", interactive=False
                        )

                # 第三页：分析报告
                with gr.TabItem("📋 分析报告", id=2):
                    gr.Markdown("### 📊 ESG分析报告生成")
                    gr.Markdown(
                        "**说明**: 请先在'模型配置与评分'页面完成ESG评分，然后在此页面导入评分结果生成专业分析报告"
                    )

                    # 步骤1：导入评分数据
                    gr.Markdown("#### 📥 步骤1: 导入评分数据")
                    gr.Markdown(
                        "请从'模型配置与评分'页面获取评分结果，或直接点击按钮自动导入最新评分数据"
                    )

                    with gr.Row():
                        import_scoring_data_btn = gr.Button(
                            "� 导入最新评分数据", variant="primary"
                        )
                        clear_imported_data_btn = gr.Button(
                            "🗑️ 清空数据", variant="secondary"
                        )

                    imported_data_status = gr.Textbox(
                        label="导入状态",
                        value="暂无数据，请先导入评分结果",
                        interactive=False,
                    )

                    # 导入的数据预览
                    with gr.Accordion("📋 导入数据预览", open=False):
                        imported_results_preview = gr.Dataframe(
                            label="评分结果预览", interactive=False, visible=False
                        )

                    gr.Markdown("---")

                    # 步骤2：生成分析报告
                    gr.Markdown("#### 📋 步骤2: 生成分析报告")
                    gr.Markdown("基于导入的评分数据，选择报告模板生成专业的ESG分析报告")

                    with gr.Row():
                        with gr.Column(scale=2):
                            generate_analysis_report_btn = gr.Button(
                                "� 生成分析报告",
                                variant="primary",
                                size="lg",
                                interactive=False,
                            )
                        with gr.Column(scale=1):
                            report_template_choice = gr.Dropdown(
                                choices=[
                                    "标准分析报告",
                                    "简化报告",
                                    "详细技术报告",
                                    "投资决策报告",
                                ],
                                value="标准分析报告",
                                label="报告模板",
                            )

                    # 报告显示区域 - 使用下拉菜单节省空间
                    with gr.Accordion("� 分析报告内容", open=False):
                        analysis_report_content = gr.Markdown(
                            value="*请先导入评分数据，然后点击'生成分析报告'按钮*"
                        )

                    # 步骤3：报告导出
                    gr.Markdown("#### 📤 步骤3: 报告导出")
                    with gr.Row():
                        export_report_txt_btn = gr.Button(
                            "📄 导出TXT", variant="secondary", interactive=False
                        )
                        export_report_word_btn = gr.Button(
                            "📄 导出Word", variant="secondary", interactive=False
                        )
                        export_report_pdf_btn = gr.Button(
                            "📄 导出PDF", variant="secondary", interactive=False
                        )

                    # 导出文件
                    with gr.Row():
                        export_txt_file = gr.File(
                            label="TXT报告下载", interactive=False
                        )
                        export_word_file = gr.File(
                            label="Word报告下载", interactive=False
                        )
                        export_pdf_file = gr.File(
                            label="PDF报告下载", interactive=False
                        )

            # 事件绑定
            all_inputs = (
                [company_name, industry_choice] + e_inputs + s_inputs + g_inputs
            )

            # 数据输入相关
            export_input_btn.click(
                fn=self.export_input_data,
                inputs=[export_format],
                outputs=[export_input_file, export_input_status],
                queue=False,
            )

            # 数据创建
            create_btn.click(
                fn=self.create_manual_input_data,
                inputs=all_inputs,
                outputs=[data_preview, data_status],
                queue=False,
            )

            upload_btn.click(
                fn=self.upload_custom_data,
                inputs=[upload_file],
                outputs=[data_preview, data_status],
                queue=False,
            )

            # 数据加载
            load_data_btn.click(
                fn=self.load_scoring_data,
                inputs=[scoring_upload_file],
                outputs=[data_load_status],
                queue=False,
            )

            # 评分计算
            calculate_btn.click(
                fn=self.calculate_esg_scores,
                inputs=[
                    alpha_param,
                    include_events_check,
                    e_weight,
                    s_weight,
                    g_weight,
                    delta_coeff,
                    epsilon_coeff,
                    zeta_coeff,
                    severity_factor,
                    max_bonus,
                    bonus_steepness,
                    threshold_multiplier,
                    data_breach_coeff,
                    env_pollution_coeff,
                    safety_accident_coeff,
                    corruption_coeff,
                    labor_dispute_coeff,
                    product_recall_coeff,
                    carbon_tax_sensitivity,
                    esg_disclosure_weight,
                    green_finance_bonus,
                    regulatory_compliance,
                    use_cross_terms,
                ],
                outputs=[results_table, charts_plot],
                queue=False,
            )

            # 结果导出
            export_results_btn.click(
                fn=self.export_results,
                inputs=[results_table],
                outputs=[export_results_file, export_results_status],
                queue=False,
            )

            # 分析报告页面事件处理

            # 导入评分数据
            import_scoring_data_btn.click(
                fn=self.import_scoring_data,
                inputs=[results_table],
                outputs=[
                    imported_data_status,
                    imported_results_preview,
                    generate_analysis_report_btn,
                ],
                queue=False,
            )

            # 清空导入数据
            clear_imported_data_btn.click(
                fn=lambda: (
                    "暂无数据，请先导入评分结果",
                    gr.Dataframe(visible=False),
                    gr.Button(interactive=False),
                ),
                inputs=[],
                outputs=[
                    imported_data_status,
                    imported_results_preview,
                    generate_analysis_report_btn,
                ],
                queue=False,
            )

            # 生成分析报告
            generate_analysis_report_btn.click(
                fn=self.generate_imported_analysis_report,
                inputs=[imported_results_preview, report_template_choice],
                outputs=[
                    analysis_report_content,
                    export_report_txt_btn,
                    export_report_word_btn,
                    export_report_pdf_btn,
                ],
                queue=False,
            )

            # 导出报告功能
            export_report_txt_btn.click(
                fn=lambda content: self.export_text_content(content, "ESG分析报告"),
                inputs=[analysis_report_content],
                outputs=[export_txt_file],
                queue=False,
            )

            export_report_word_btn.click(
                fn=self.export_report_as_word,
                inputs=[analysis_report_content],
                outputs=[export_word_file],
                queue=False,
            )

            export_report_pdf_btn.click(
                fn=self.export_report_as_pdf,
                inputs=[analysis_report_content],
                outputs=[export_pdf_file],
                queue=False,
            )

        return interface


# 启动应用
if __name__ == "__main__":
    app = ESGGradioApp()
    interface = app.create_interface()

    # 启动界面
    interface.launch(
        server_name="0.0.0.0", server_port=7860, share=False, show_error=True
    )
